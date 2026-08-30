import uuid
import base64
import struct
import zlib
import pytest
from io import BytesIO
from zipfile import ZipFile

from docx import Document
from docx.shared import RGBColor
from fastapi.testclient import TestClient
from fontTools.ttLib import TTFont
from lxml import etree

from app.main import app


client = TestClient(app)


@pytest.mark.parametrize("span", [1, 2])
def test_large_table_image_fits_merged_cell(span: int) -> None:
    def chunk(tag, content):
        return struct.pack(">I", len(content)) + tag + content + struct.pack(">I", zlib.crc32(tag + content))
    png = b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", struct.pack(">IIBBBBB", 1200, 240, 8, 2, 0, 0, 0))
    png += chunk(b"IDAT", zlib.compress((b"\0" + b"\0\0\xff" * 1200) * 240)) + chunk(b"IEND", b"")
    image = base64.b64encode(png).decode()
    html = f"<table><tr><td colspan='{span}'><img src='data:image/png;base64,{image}'></td><td>" + "实验记录" * 100 + "</td></tr></table>"
    response = client.post("/api/documents/export-docx", json={"html": html})
    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    width = document.tables[0].cell(0, 0).width
    shape = document.inline_shapes[0]
    assert shape.width <= width - 280 * 635
    assert abs(shape.width - shape.height * 5) <= 5


def test_narrative_table_widths_and_top_alignment_are_content_based() -> None:
    row = "<tr><td>" + "观测记录" * 150 + "</td><td>" + "实验结论" * 20 + "</td><td>备注</td></tr>"
    html = "<table><tr><th>记录</th><th>结果</th><th>状态</th></tr>" + row + "</table>"
    banner = "<tr><td colspan='3'>" + "通栏说明" * 200 + "</td></tr>"
    documents = []
    for source in (html, html.replace(row, banner + row)):
        response = client.post("/api/documents/export-docx", json={"html": source})
        assert response.status_code == 200
        documents.append(Document(BytesIO(response.content)))
    table = documents[0].tables[0]
    widths = [int(value) for value in table._tbl.xpath("./w:tblGrid/w:gridCol/@w:w")]
    assert widths[0] > widths[1] > widths[2] >= 720
    assert sum(widths) == 9360
    assert widths == [int(value) for value in documents[1].tables[0]._tbl.xpath("./w:tblGrid/w:gridCol/@w:w")]
    assert table.cell(0, 0)._tc.xpath("./w:tcPr/w:vAlign/@w:val") == ["center"]
    assert table.cell(1, 0)._tc.xpath("./w:tcPr/w:vAlign/@w:val") == ["top"]
    assert not table._tbl.xpath(".//w:cantSplit | .//w:keepNext")


@pytest.mark.parametrize("html", [
    "<table><tr><td colspan='2147483647'>x</td></tr></table>",
    "<table><tr><td rowspan='2147483647'>x</td></tr></table>",
    "<table><tr><td colspan='60'>x</td><td colspan='60'>y</td></tr></table>",
    "<table>" + "<tr><td colspan='100'>x</td></tr>" * 101 + "</table>",
    "<table>" + "<tr><td>x</td></tr>" * 2001 + "</table>",
])
def test_export_rejects_oversized_table_before_expansion(html: str) -> None:
    response = client.post("/api/documents/export-docx", json={"html": html})
    assert response.status_code == 413
    assert response.json()["error"]["code"] == "TABLE_TOO_LARGE"


@pytest.mark.parametrize("span", ["99999999999999999999999", "invalid"])
def test_export_rejects_invalid_table_span(span: str) -> None:
    response = client.post("/api/documents/export-docx", json={
        "html": f"<table><tr><td colspan='{span}'>x</td></tr></table>",
    })
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "INVALID_TABLE_SPAN"


def _read_embedded_font(archive: ZipFile, style: str) -> TTFont:
    font_table = etree.fromstring(archive.read("word/fontTable.xml"))
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    element_name = "embedRegular" if style == "regular" else "embedBold"
    font_key = font_table.xpath(
        f"./w:font[@w:name='Noto Sans SC']/w:{element_name}/@w:fontKey",
        namespaces=namespaces,
    )[0]
    font_data = bytearray(archive.read(f"word/fonts/NotoSansSC-{style}.odttf"))
    key = uuid.UUID(font_key.strip("{}")).bytes[::-1]
    for index in range(min(32, len(font_data))):
        font_data[index] ^= key[index % len(key)]
    return TTFont(BytesIO(font_data))


def test_export_docx_converts_supported_html() -> None:
    response = client.post(
        "/api/documents/export-docx",
        json={
            "filename": "三角函数教案.docx",
            "html": """
                <article>
                  <h1 data-docx-title>三角函数教案</h1>
                  <h2>学习目标</h2>
                  <p>掌握 <strong>正弦函数</strong> 与 <em>余弦函数</em>。</p>
                  <ol><li>复习定义</li><li>完成练习</li></ol>
                  <p><span data-latex="\\sin^2 x + \\cos^2 x = 1"></span></p>
                  <table><tr><th>角度</th><th>正弦值</th></tr><tr><td>范围：<span data-latex="360^{\\circ} \\leq b &lt; 720^{\\circ}"></span> 的元素</td><td>1/2</td></tr></table>
                  <script>alert('ignored')</script>
                </article>
            """,
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "filename*=UTF-8''" in response.headers["content-disposition"]

    document = Document(BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "三角函数教案" in text
    assert "学习目标" in text
    assert "正弦函数" in text
    assert r"\sin^2 x + \cos^2 x = 1" not in text
    assert "alert" not in text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == "范围： 的元素"
    assert document.tables[0].cell(1, 1).text == "1/2"

    with ZipFile(BytesIO(response.content)) as archive:
        content_types_xml = archive.read("[Content_Types].xml").decode()
        document_xml = archive.read("word/document.xml").decode()
        font_table_xml = archive.read("word/fontTable.xml").decode()
        font_rels_xml = archive.read("word/_rels/fontTable.xml.rels").decode()
        numbering_xml = archive.read("word/numbering.xml").decode()
        settings_xml = archive.read("word/settings.xml").decode()
        styles_xml = archive.read("word/styles.xml").decode()
        embedded_font_parts = sorted(
            name for name in archive.namelist() if name.startswith("word/fonts/")
        )
        regular_font = _read_embedded_font(archive, "regular")
        bold_font = _read_embedded_font(archive, "bold")
    assert "w:numPr" in document_xml
    assert "w:tblHeader" in document_xml
    assert document_xml.count("<m:oMath>") == 2
    assert "<m:sSup>" in document_xml
    assert "<m:t>≤</m:t>" in document_xml
    assert "<m:t>∘</m:t>" in document_xml
    assert "\\sin^2" not in document_xml
    assert "360^{" not in document_xml
    assert "multilevel" in numbering_xml
    assert 'w:eastAsia="zh-CN"' in styles_xml
    assert 'w:eastAsia="Noto Sans SC"' in styles_xml
    assert embedded_font_parts == [
        "word/fonts/NotoSansSC-bold.odttf",
        "word/fonts/NotoSansSC-regular.odttf",
    ]
    assert 'w:name="Noto Sans SC"' in font_table_xml
    assert 'w:altName w:val="Microsoft YaHei"' in font_table_xml
    assert "w:embedRegular" in font_table_xml
    assert "w:embedBold" in font_table_xml
    assert font_rels_xml.count("/relationships/font") == 2
    assert 'Extension="odttf"' in content_types_xml
    assert "w:embedTrueTypeFonts" in settings_xml
    assert "w:saveSubsetFonts" in settings_xml
    assert ord("三") in regular_font.getBestCmap()
    assert ord("教") in bold_font.getBestCmap()
    assert regular_font["OS/2"].fsType == 0
    assert bold_font["OS/2"].fsType == 0


def test_export_docx_keeps_invalid_latex_visible() -> None:
    response = client.post(
        "/api/documents/export-docx",
        json={"filename": "invalid-formula.docx", "html": '<p><span data-latex="\\frac{"></span></p>'},
    )

    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    assert document.paragraphs[0].text == r"\frac{"


def test_export_docx_rejects_html_without_content() -> None:
    response = client.post(
        "/api/documents/export-docx",
        json={"filename": "empty.docx", "html": "<script>alert(1)</script>"},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "EMPTY_HTML"


def test_export_docx_embeds_data_url_image() -> None:
    response = client.post(
        "/api/documents/export-docx",
        json={
            "filename": "image.docx",
            "html": (
                '<p>Image</p><img alt="pixel" '
                'src="data:image/png;base64,'
                'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="/>'
            ),
        },
    )

    assert response.status_code == 200
    assert len(Document(BytesIO(response.content)).inline_shapes) == 1


def test_export_docx_distributes_colspan_content_across_merged_columns() -> None:
    response = client.post(
        "/api/documents/export-docx",
        json={
            "filename": "layout-table.docx",
            "html": """
                <table class="lesson-layout-table" data-repeat-header="false">
                  <tr><td colspan="4"><strong>一、内容</strong></td></tr>
                  <tr><td>A</td><td>B</td><td colspan="2">这是需要在合并单元格中获得足够宽度显示的子问题内容，列宽计算不能忽略跨列单元格中的长文本内容</td></tr>
                </table>
            """,
        },
    )

    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    assert len(document.tables) == 1
    with ZipFile(BytesIO(response.content)) as archive:
        document_xml = archive.read("word/document.xml").decode()
    assert "w:tblHeader" not in document_xml

    grid = document.tables[0]._tbl.tblGrid
    widths = [int(column.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) for column in grid]
    assert widths[2] > widths[0]
    assert widths[3] > widths[1]


def test_export_docx_preserves_lesson_heading_layout_and_skips_preview_title() -> None:
    response = client.post(
        "/api/documents/export-docx",
        json={
            "filename": "lesson-title-layout.docx",
            "html": """
                <article>
                  <header class="document-title"><h1>预览元数据标题</h1></header>
                  <div class="document-block document-block-heading-center">
                    <h1 class="lesson-heading lesson-heading-1 lesson-heading-center">第一课时</h1>
                  </div>
                  <div class="document-block document-block-heading-center">
                    <h2 class="lesson-heading lesson-heading-2 lesson-heading-center">5.1.1　　任意角</h2>
                  </div>
                  <table class="lesson-layout-table"><tr><td>一、内容</td></tr></table>
                </article>
            """,
        },
    )

    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "第一课时",
        "5.1.1　　任意角",
    ]
    assert [paragraph.style.name for paragraph in document.paragraphs] == [
        "Lesson Heading 1",
        "Lesson Heading 2",
    ]
    assert all(
        paragraph.alignment == 1
        for paragraph in document.paragraphs
    )
    assert document.paragraphs[0].style.font.color.rgb == RGBColor(0, 0, 0)
    assert document.paragraphs[1].style.font.color.rgb == RGBColor(0, 0, 0)
