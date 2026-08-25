import base64
import binascii
import re
from io import BytesIO
from pathlib import Path
from urllib.parse import unquote_to_bytes

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml

from app.core.exceptions import AppError
from app.services.docx_font_embedder import FONT_FAMILY, embed_default_fonts

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 120, "bottom": 120, "start": 140, "end": 140}
EAST_ASIA_FONT = FONT_FAMILY


class DocxExportService:
    """Convert sanitized, semantic HTML into a Word document."""

    def export(self, html: str, filename: str) -> tuple[BytesIO, str]:
        soup = BeautifulSoup(html, "html.parser")
        for unsafe in soup.select("script, style, noscript, iframe, object, embed"):
            unsafe.decompose()

        document = Document()
        self._configure_document(document)
        bullet_num_id = self._add_numbering(document, ordered=False)
        decimal_num_id = self._add_numbering(document, ordered=True)
        root = soup.body or soup
        state = {"has_content": False, "bullet_num_id": bullet_num_id, "decimal_num_id": decimal_num_id}
        for child in list(root.children):
            self._append_block(document, child, state)

        if not state["has_content"]:
            raise AppError("EMPTY_HTML", "HTML 中没有可导出的内容", 422)

        output = BytesIO()
        document.save(output)
        characters = "".join(element.text or "" for element in document.element.iter())
        return embed_default_fonts(output, characters), self._safe_filename(filename)

    def _configure_document(self, document: DocumentObject) -> None:
        self._set_document_language(document, "zh-CN")
        section = document.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        self._set_style(document, "Normal", size=11, color="000000", before=0, after=6, line=1.1, bold=False)
        title_style = document.styles.add_style("HTML Title", WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = document.styles["Normal"]
        self._set_style(document, "HTML Title", size=24, color="0B2545", before=0, after=12, line=1.0, bold=True)
        self._set_style(document, "Heading 1", size=16, color="2E74B5", before=16, after=8, line=1.0, bold=True)
        self._set_style(document, "Heading 2", size=13, color="2E74B5", before=12, after=6, line=1.0, bold=True)
        self._set_style(document, "Heading 3", size=12, color="1F4D78", before=8, after=4, line=1.0, bold=True)
        self._set_style(document, "Caption", size=9, color="68736E", before=0, after=6, line=1.0, bold=False)

    @staticmethod
    def _set_document_language(document: DocumentObject, language: str) -> None:
        defaults = document.styles.element.find(qn("w:docDefaults"))
        if defaults is None:
            defaults = OxmlElement("w:docDefaults")
            document.styles.element.insert(0, defaults)
        run_defaults = defaults.find(qn("w:rPrDefault"))
        if run_defaults is None:
            run_defaults = OxmlElement("w:rPrDefault")
            defaults.append(run_defaults)
        run_properties = run_defaults.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            run_defaults.append(run_properties)
        fonts = run_properties.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            run_properties.insert(0, fonts)
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), EAST_ASIA_FONT)
        language_element = run_properties.find(qn("w:lang"))
        if language_element is None:
            language_element = OxmlElement("w:lang")
            run_properties.append(language_element)
        language_element.set(qn("w:eastAsia"), language)
        theme_language = document.settings.element.find(qn("w:themeFontLang"))
        if theme_language is not None:
            theme_language.set(qn("w:eastAsia"), language)

    @staticmethod
    def _set_style(
        document: DocumentObject,
        name: str,
        *,
        size: int,
        color: str,
        before: int,
        after: int,
        line: float,
        bold: bool,
    ) -> None:
        style = document.styles[name]
        style.font.name = EAST_ASIA_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        fonts = style._element.get_or_add_rPr().rFonts
        fonts.set(qn("w:ascii"), EAST_ASIA_FONT)
        fonts.set(qn("w:hAnsi"), EAST_ASIA_FONT)
        fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line

    def _append_block(self, document: DocumentObject, node: object, state: dict[str, object]) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                document.add_paragraph(text)
                state["has_content"] = True
            return
        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name in {"script", "style", "noscript", "iframe", "object", "embed", "button", "input", "textarea"}:
            return
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(name[1])
            is_title = node.has_attr("data-docx-title") or (level == 1 and not state["has_content"])
            paragraph = document.add_paragraph(style="HTML Title" if is_title else f"Heading {min(level, 3)}")
            self._append_inline(paragraph, node)
            state["has_content"] = True
            return
        if name in {"p", "pre", "blockquote"}:
            paragraph = document.add_paragraph()
            if name == "blockquote":
                paragraph.paragraph_format.left_indent = Inches(0.3)
                paragraph.paragraph_format.right_indent = Inches(0.3)
            if name == "pre":
                run = paragraph.add_run(node.get_text("\n", strip=False))
                self._set_run_font(run, "Courier New", EAST_ASIA_FONT)
                run.font.size = Pt(9.5)
            else:
                self._append_inline(paragraph, node)
            if paragraph.text.strip() or paragraph._p.xpath(".//w:drawing | .//m:oMath"):
                state["has_content"] = True
            else:
                paragraph._element.getparent().remove(paragraph._element)
            return
        if name in {"ul", "ol"}:
            num_id = int(state["decimal_num_id"] if name == "ol" else state["bullet_num_id"])
            self._append_list(document, node, num_id, 0)
            state["has_content"] = True
            return
        if name == "table":
            if self._append_table(document, node):
                state["has_content"] = True
            return
        if name == "figure":
            image = node.find("img")
            if image and self._append_standalone_image(document, image):
                state["has_content"] = True
            caption = node.find("figcaption")
            if caption and caption.get_text(" ", strip=True):
                paragraph = document.add_paragraph(style="Caption")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._append_inline(paragraph, caption)
            return
        if name == "img":
            if self._append_standalone_image(document, node):
                state["has_content"] = True
            return
        if node.has_attr("data-latex"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._append_formula(paragraph, str(node["data-latex"]))
            state["has_content"] = True
            return

        for child in list(node.children):
            self._append_block(document, child, state)

    def _append_inline(
        self,
        paragraph,
        node: object,
        *,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
    ) -> None:
        if isinstance(node, NavigableString):
            text = re.sub(r"[\t\r\n ]+", " ", str(node))
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
            return
        if not isinstance(node, Tag):
            return
        if node.has_attr("data-latex"):
            self._append_formula(paragraph, str(node["data-latex"]))
            return
        name = node.name.lower()
        if name == "br":
            paragraph.add_run().add_break()
            return
        if name == "img":
            self._append_image_to_paragraph(paragraph, node)
            return
        if name in {"script", "style", "button", "input", "textarea", "svg"}:
            return
        next_bold = bold or name in {"strong", "b", "th"}
        next_italic = italic or name in {"em", "i", "cite"}
        next_underline = underline or name in {"u", "a"}
        for child in list(node.children):
            before = len(paragraph.runs)
            self._append_inline(
                paragraph,
                child,
                bold=next_bold,
                italic=next_italic,
                underline=next_underline,
            )
            for run in paragraph.runs[before:]:
                if name == "sup":
                    run.font.superscript = True
                elif name == "sub":
                    run.font.subscript = True
                if name == "a":
                    run.font.color.rgb = RGBColor(0x05, 0x66, 0x5C)

    def _append_formula(self, paragraph, latex: str) -> None:
        value = latex.strip()
        if not value:
            return

        try:
            mathml = latex_to_mathml(value)
            omml = mathml_to_omml(mathml)
            formula = parse_xml(omml.replace("<m:oMath>", f"<m:oMath {nsdecls('m')}>", 1))
        except Exception:
            # Keep malformed OCR output visible instead of dropping content from
            # the exported document. Valid LaTeX is emitted as editable OMML.
            run = paragraph.add_run(value)
            self._set_run_font(run, "Cambria Math", EAST_ASIA_FONT)
            return

        paragraph._p.append(formula)

    def _append_list(self, document: DocumentObject, element: Tag, num_id: int, level: int) -> None:
        for item in element.find_all("li", recursive=False):
            paragraph = document.add_paragraph()
            self._set_numbering(paragraph, num_id, min(level, 8))
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
            for child in list(item.children):
                if isinstance(child, Tag) and child.name.lower() in {"ul", "ol"}:
                    continue
                self._append_inline(paragraph, child)
            for nested in item.find_all(["ul", "ol"], recursive=False):
                self._append_list(document, nested, num_id, level + 1)

    def _append_table(self, document: DocumentObject, element: Tag) -> bool:
        source_rows = element.find_all("tr")
        if not source_rows:
            return False

        placements: list[tuple[int, int, int, int, Tag]] = []
        occupied: set[tuple[int, int]] = set()
        max_columns = 0
        for row_index, source_row in enumerate(source_rows):
            column = 0
            for source_cell in source_row.find_all(["th", "td"], recursive=False):
                while (row_index, column) in occupied:
                    column += 1
                rowspan = self._positive_int(source_cell.get("rowspan"), 1)
                colspan = self._positive_int(source_cell.get("colspan"), 1)
                placements.append((row_index, column, rowspan, colspan, source_cell))
                for r in range(row_index, row_index + rowspan):
                    for c in range(column, column + colspan):
                        occupied.add((r, c))
                column += colspan
            max_columns = max(max_columns, column)
        if max_columns == 0:
            return False

        table = document.add_table(rows=len(source_rows), cols=max_columns)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        self._set_table_properties(table)
        widths = self._table_widths(placements, max_columns)
        self._set_table_grid(table, widths)

        for row_index, column, rowspan, colspan, source_cell in placements:
            end_row = min(row_index + rowspan - 1, len(source_rows) - 1)
            end_col = min(column + colspan - 1, max_columns - 1)
            cell = table.cell(row_index, column)
            if end_row != row_index or end_col != column:
                cell = cell.merge(table.cell(end_row, end_col))
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.2
            self._append_inline(paragraph, source_cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            self._set_cell_width(cell, sum(widths[column : end_col + 1]))
            classes = set(source_cell.get("class") or [])
            if "lesson-layout-centered-cell" in classes:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if source_cell.name.lower() == "th":
                self._shade_cell(cell, "F2F4F7")
                for run in paragraph.runs:
                    run.bold = True

        first_source_row = source_rows[0]
        if (
            str(element.get("data-repeat-header", "true")).lower() != "false"
            and first_source_row.find("th", recursive=False) is not None
        ):
            self._repeat_table_header(table.rows[0])
        return True

    @staticmethod
    def _table_widths(placements: list[tuple[int, int, int, int, Tag]], count: int) -> list[int]:
        weights = [8] * count
        for _row, column, _rowspan, colspan, cell in placements:
            content_weight = min(len(cell.get_text(" ", strip=True)), 40 * colspan)
            per_column_weight = max(8, round(content_weight / colspan))
            for index in range(column, min(column + colspan, count)):
                weights[index] = max(weights[index], per_column_weight)
        total = sum(weights)
        widths = [max(720, round(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
        scale = CONTENT_WIDTH_DXA / sum(widths)
        widths = [round(width * scale) for width in widths]
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
        return widths

    @staticmethod
    def _set_table_properties(table) -> None:
        properties = table._tbl.tblPr
        width = properties.first_child_found_in("w:tblW")
        if width is None:
            width = OxmlElement("w:tblW")
            properties.append(width)
        width.set(qn("w:type"), "dxa")
        width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
        indent = OxmlElement("w:tblInd")
        indent.set(qn("w:type"), "dxa")
        indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
        properties.append(indent)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        properties.append(layout)
        margins = OxmlElement("w:tblCellMar")
        for side, value in CELL_MARGIN_DXA.items():
            margin = OxmlElement(f"w:{side}")
            margin.set(qn("w:w"), str(value))
            margin.set(qn("w:type"), "dxa")
            margins.append(margin)
        properties.append(margins)

    @staticmethod
    def _set_table_grid(table, widths: list[int]) -> None:
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            column = OxmlElement("w:gridCol")
            column.set(qn("w:w"), str(width))
            grid.append(column)
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                DocxExportService._set_cell_width(cell, widths[min(index, len(widths) - 1)])

    @staticmethod
    def _set_cell_width(cell, width: int) -> None:
        tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
        tc_width.set(qn("w:type"), "dxa")
        tc_width.set(qn("w:w"), str(width))

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            cell._tc.get_or_add_tcPr().append(shading)
        shading.set(qn("w:fill"), fill)

    @staticmethod
    def _repeat_table_header(row) -> None:
        row_properties = row._tr.get_or_add_trPr()
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        row_properties.append(header)

    def _append_standalone_image(self, document: DocumentObject, image: Tag) -> bool:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self._append_image_to_paragraph(paragraph, image):
            return True
        paragraph._element.getparent().remove(paragraph._element)
        return False

    def _append_image_to_paragraph(self, paragraph, image: Tag) -> bool:
        payload = self._decode_data_image(str(image.get("src", "")))
        if payload is None:
            alt = str(image.get("alt", "")).strip()
            if alt:
                run = paragraph.add_run(f"[图片：{alt}]")
                run.italic = True
                return True
            return False
        try:
            shape = paragraph.add_run().add_picture(BytesIO(payload))
        except (ValueError, OSError, TypeError):
            return False
        max_width = Inches(6.25)
        max_height = Inches(7.5)
        scale = min(1.0, max_width / shape.width, max_height / shape.height)
        shape.width = int(shape.width * scale)
        shape.height = int(shape.height * scale)
        alt = str(image.get("alt", "")).strip()
        if alt:
            shape._inline.docPr.set("descr", alt)
        return True

    @staticmethod
    def _decode_data_image(src: str) -> bytes | None:
        match = re.fullmatch(r"data:image/(png|jpeg|jpg|gif);(base64)?,(.+)", src, re.IGNORECASE | re.DOTALL)
        if not match:
            return None
        try:
            return base64.b64decode(match.group(3), validate=True) if match.group(2) else unquote_to_bytes(match.group(3))
        except (binascii.Error, ValueError):
            return None

    @staticmethod
    def _add_numbering(document: DocumentObject, *, ordered: bool) -> int:
        numbering = document.part.numbering_part.element
        abstract_ids = [int(value) for value in numbering.xpath("./w:abstractNum/@w:abstractNumId")]
        num_ids = [int(value) for value in numbering.xpath("./w:num/@w:numId")]
        abstract_id = max(abstract_ids, default=0) + 1
        num_id = max(num_ids, default=0) + 1

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi_level = OxmlElement("w:multiLevelType")
        multi_level.set(qn("w:val"), "multilevel")
        abstract.append(multi_level)
        for level in range(9):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
            lvl.append(num_fmt)
            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), f"%{level + 1}." if ordered else "•")
            lvl.append(lvl_text)
            suffix = OxmlElement("w:suff")
            suffix.set(qn("w:val"), "tab")
            lvl.append(suffix)
            ppr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(720 + level * 360))
            tabs.append(tab)
            ppr.append(tabs)
            indent = OxmlElement("w:ind")
            indent.set(qn("w:left"), str(720 + level * 360))
            indent.set(qn("w:hanging"), "360")
            ppr.append(indent)
            lvl.append(ppr)
            abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    @staticmethod
    def _set_numbering(paragraph, num_id: int, level: int) -> None:
        properties = paragraph._p.get_or_add_pPr()
        num_properties = properties.get_or_add_numPr()
        num_properties.get_or_add_ilvl().val = level
        num_properties.get_or_add_numId().val = num_id

    @staticmethod
    def _set_run_font(run, western: str, east_asia: str) -> None:
        run.font.name = western
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), western)
        fonts.set(qn("w:hAnsi"), western)
        fonts.set(qn("w:eastAsia"), east_asia)

    @staticmethod
    def _positive_int(value: object, default: int) -> int:
        try:
            return max(1, int(str(value)))
        except (TypeError, ValueError):
            return default

    @staticmethod
    def _safe_filename(filename: str) -> str:
        name = Path(filename.replace("\\", "/")).name.strip()
        name = re.sub(r"[\x00-\x1f<>:\"/\\|?*]", "_", name)
        stem = Path(name).stem.strip(" .") or "lesson"
        return f"{stem[:160]}.docx"
